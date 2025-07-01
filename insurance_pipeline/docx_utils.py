import re
from docx import Document
from collections import defaultdict
import convertapi
from insurance_pipeline.config import CONVERTAPI_API_KEY

convertapi.api_credentials = CONVERTAPI_API_KEY

def extract_field_contexts(doc_path):
    """
    Returns a dict mapping field -> list of sentences containing that placeholder.
    """
    doc = Document(doc_path)
    pattern = r"\[([A-Z0-9_]+)\]"
    contexts = defaultdict(list)

    for para in doc.paragraphs:
        sent = para.text.strip()
        matches = re.findall(pattern, sent)
        for match in matches:
            contexts[match].append(sent)

    return dict(contexts)

def fill_placeholders_in_docx(template_path, extracted_values):
    doc = Document(template_path)
    pattern = r"\[([A-Z0-9_]+)\]"

    def replace_in_text(text):
        return re.sub(pattern, lambda m: extracted_values.get(m.group(1), m.group(0)) or m.group(0), text)

    for para in doc.paragraphs:
        para.text = replace_in_text(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace_in_text(cell.text)
                
    return doc

def convert_docx_to_pdf_via_convertapi(input_path: str, output_dir: str) -> str:
    """Converts a .docx to .pdf using ConvertAPI."""
    result = convertapi.convert(
        'pdf',
        {'File': input_path},
        from_format='docx'
    )
    saved_paths = result.save_files(output_dir)
    return saved_paths[0]